"""Tests for vendor document (PDF/DOCX/Markdown) conversion and segmentation."""

from __future__ import annotations

from collections.abc import Callable
from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from ingestion_pipeline.documents import (
    DocumentUnit,
    convert_to_markdown,
    docx_to_markdown,
    markdown_object_name,
    original_filename_from_markdown_object,
    pdf_to_markdown,
    split_markdown_units,
    supported_extensions,
)


def _docx_bytes(build: Callable[[Document], None]) -> bytes:
    document = Document()
    build(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestPdfToMarkdown:
    @patch("ingestion_pipeline.documents.PdfReader")
    def test_renders_one_heading_per_non_empty_page(self, mock_reader_cls):
        page1 = MagicMock()
        page1.extract_text.return_value = "Page one text"
        page2 = MagicMock()
        page2.extract_text.return_value = "Page two text"
        mock_reader_cls.return_value.pages = [page1, page2]

        markdown = pdf_to_markdown(b"fake-pdf-bytes")

        assert markdown == "## Page 1\n\nPage one text\n\n## Page 2\n\nPage two text"

    @patch("ingestion_pipeline.documents.PdfReader")
    def test_skips_empty_and_none_pages(self, mock_reader_cls):
        empty_page = MagicMock()
        empty_page.extract_text.return_value = "   "
        text_page = MagicMock()
        text_page.extract_text.return_value = "Has content"
        none_page = MagicMock()
        none_page.extract_text.return_value = None

        mock_reader_cls.return_value.pages = [empty_page, text_page, none_page]

        markdown = pdf_to_markdown(b"fake-pdf-bytes")

        assert markdown == "## Page 2\n\nHas content"

    @patch("ingestion_pipeline.documents.PdfReader")
    def test_no_pages_yields_empty_string(self, mock_reader_cls):
        mock_reader_cls.return_value.pages = []

        assert pdf_to_markdown(b"fake-pdf-bytes") == ""

    @patch("ingestion_pipeline.documents.PdfReader")
    def test_escapes_leading_hashes_in_page_text(self, mock_reader_cls):
        page = MagicMock()
        page.extract_text.return_value = "## Not a real heading\nRegular line"
        mock_reader_cls.return_value.pages = [page]

        markdown = pdf_to_markdown(b"fake-pdf-bytes")

        assert "\\## Not a real heading" in markdown
        # The escaped line must not be picked up as a section boundary downstream.
        units = split_markdown_units(markdown)
        assert len(units) == 1
        assert units[0].attributes == {"section": "Page 1"}


class TestDocxToMarkdown:
    def test_maps_heading_styles_to_atx_prefixes(self):
        def build(document):
            document.add_paragraph("Intro text before any heading")
            document.add_heading("Section One", level=1)
            document.add_paragraph("Body of section one")
            document.add_heading("Section Two", level=2)
            document.add_paragraph("Body of section two")

        markdown = docx_to_markdown(_docx_bytes(build))

        assert markdown == (
            "Intro text before any heading\n\n"
            "# Section One\n\n"
            "Body of section one\n\n"
            "## Section Two\n\n"
            "Body of section two"
        )

    def test_escapes_leading_hashes_in_body_paragraphs(self):
        def build(document):
            document.add_paragraph("# Looks like a heading but is not")

        markdown = docx_to_markdown(_docx_bytes(build))

        assert markdown == "\\# Looks like a heading but is not"

    def test_empty_document_yields_empty_string(self):
        assert docx_to_markdown(_docx_bytes(lambda document: None)) == ""


class TestSplitMarkdownUnits:
    def test_splits_on_headings_of_any_level(self):
        markdown = "Intro\n\n## Page 1\n\nFirst page body\n\n## Page 2\n\nSecond page body"

        units = split_markdown_units(markdown)

        assert units == [
            DocumentUnit(text="Intro", attributes={}),
            DocumentUnit(text="First page body", attributes={"section": "Page 1"}),
            DocumentUnit(text="Second page body", attributes={"section": "Page 2"}),
        ]

    def test_falls_back_to_single_unit_when_no_headings(self):
        markdown = "First paragraph.\n\nSecond paragraph."

        units = split_markdown_units(markdown)

        assert units == [DocumentUnit(text="First paragraph.\n\nSecond paragraph.", attributes={})]

    def test_ignores_headings_with_no_body_text(self):
        markdown = "# Empty Section\n\n# Section With Body\n\nSome content"

        units = split_markdown_units(markdown)

        assert units == [DocumentUnit(text="Some content", attributes={"section": "Section With Body"})]

    def test_empty_markdown_yields_no_units(self):
        assert split_markdown_units("") == []

    def test_recognizes_all_heading_levels(self):
        markdown = "###### Deep Heading\n\nDeep content"

        units = split_markdown_units(markdown)

        assert units == [DocumentUnit(text="Deep content", attributes={"section": "Deep Heading"})]


class TestConvertToMarkdown:
    @patch("ingestion_pipeline.documents.pdf_to_markdown")
    def test_dispatches_pdf_by_extension(self, mock_convert):
        mock_convert.return_value = "## Page 1\n\ntext"

        result = convert_to_markdown("gnodeb.pdf", b"data")

        mock_convert.assert_called_once_with(b"data")
        assert result == mock_convert.return_value

    @patch("ingestion_pipeline.documents.docx_to_markdown")
    def test_dispatches_docx_by_extension_case_insensitive(self, mock_convert):
        mock_convert.return_value = "# Heading\n\ntext"

        result = convert_to_markdown("doc.DOCX", b"data")

        mock_convert.assert_called_once_with(b"data")
        assert result == mock_convert.return_value

    def test_markdown_files_pass_through_unchanged(self):
        result = convert_to_markdown("notes.md", b"# Already markdown")

        assert result == "# Already markdown"

    def test_raises_for_unsupported_extension(self):
        with pytest.raises(ValueError, match="Unsupported vendor document type"):
            convert_to_markdown("notes.txt", b"data")

    def test_raises_for_missing_extension(self):
        with pytest.raises(ValueError, match="Unsupported vendor document type"):
            convert_to_markdown("README", b"data")


class TestSupportedExtensions:
    def test_includes_pdf_docx_and_markdown(self):
        assert supported_extensions() == frozenset({".pdf", ".docx", ".md"})


class TestMarkdownObjectNaming:
    def test_markdown_object_name_appends_md_suffix(self):
        assert markdown_object_name("gnodeb.pdf") == "gnodeb.pdf.md"
        assert markdown_object_name("ran_metrics_and_anomalies.docx") == "ran_metrics_and_anomalies.docx.md"

    def test_original_filename_from_markdown_object_round_trips(self):
        for filename in ["gnodeb.pdf", "ran_metrics_and_anomalies.docx", "notes.md"]:
            object_name = markdown_object_name(filename)
            assert original_filename_from_markdown_object(object_name) == filename

    def test_original_filename_from_markdown_object_preserves_prefix(self):
        assert original_filename_from_markdown_object("telco-docs/gnodeb.pdf.md") == "telco-docs/gnodeb.pdf"
